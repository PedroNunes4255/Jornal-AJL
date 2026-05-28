# gerar_noticia_v2.py

import os
import re
import json
import shutil
import unicodedata

from pathlib import Path
from datetime import datetime
from docx import Document

# ==========================================
# CONFIG
# ==========================================

PASTA_FRONTEND = "./frontend"

PASTA_PAGES = os.path.join(
    PASTA_FRONTEND,
    "pages"
)

PASTA_JS = os.path.join(
    PASTA_FRONTEND,
    "js"
)

PASTA_IMG = os.path.join(
    PASTA_FRONTEND,
    "IMG"
)

NEWS_DATA_ARQUIVO = os.path.join(
    PASTA_JS,
    "news-data.js"
)

AUTOR_PADRAO = "@jornalajl"

IMAGEM_PADRAO = "default.jpg"

CATEGORIAS = {
    "1": "Cultura",
    "2": "Brasil",
    "3": "Escola",
    "4": "Esportes",
    "5": "Mundo"
}

# ==========================================
# UTILIDADES
# ==========================================

def limpar_acentos(texto):

    return ''.join(
        c for c in unicodedata.normalize('NFD', texto)
        if unicodedata.category(c) != 'Mn'
    )

def slugify(texto):

    texto = limpar_acentos(
        texto.lower()
    )

    texto = re.sub(
        r'[^a-z0-9\s-]',
        '',
        texto
    )

    texto = re.sub(
        r'[\s-]+',
        '-',
        texto
    )

    return texto.strip('-')

# ==========================================
# EXCERPT MELHORADO
# ==========================================

def gerar_excerpt(texto, tamanho=260):

    linhas_ruins = [
        "foto:",
        "imagem:",
        "crédito:",
        "créditos:",
        "getty",
        "reprodução",
        "fonte:",
        "nurphoto",
        "via",
        "ap photo"
    ]

    paragrafos = texto.split("\n")

    paragrafos_filtrados = []

    for p in paragrafos:

        p_limpo = p.strip()

        if not p_limpo:
            continue

        p_lower = p_limpo.lower()

        ignorar = False

        for lixo in linhas_ruins:

            if lixo in p_lower:
                ignorar = True
                break

        if not ignorar:

            paragrafos_filtrados.append(
                p_limpo
            )

    texto_final = " ".join(
        paragrafos_filtrados
    )

    if len(texto_final) <= tamanho:
        return texto_final

    return (
        texto_final[:tamanho]
        .rsplit(' ', 1)[0]
        + "…"
    )

# ==========================================
# LER DOCX
# ==========================================

def ler_docx(caminho):

    doc = Document(caminho)

    paragrafos = []

    for p in doc.paragraphs:

        texto = p.text.strip()

        if texto:
            paragrafos.append(texto)

    return "\n".join(paragrafos)

# ==========================================
# EXTRAIR IMAGEM DO DOCX
# ==========================================

def extrair_imagem_docx(caminho_docx, slug):

    try:

        doc = Document(caminho_docx)

        rels = doc.part.rels

        contador = 1

        for rel in rels.values():

            if "image" in rel.target_ref:

                imagem = rel.target_part.blob

                nome = f"{slug}-{contador}.jpg"

                caminho_img = os.path.join(
                    PASTA_IMG,
                    nome
                )

                with open(
                    caminho_img,
                    "wb"
                ) as f:

                    f.write(imagem)

                print(
                    f"\n🖼️ Imagem extraída automaticamente: {nome}"
                )

                return nome

    except Exception as e:

        print(
            f"\n⚠️ Erro ao extrair imagem: {e}"
        )

    return IMAGEM_PADRAO

# ==========================================
# COPIAR IMAGEM MANUAL
# ==========================================

def copiar_imagem_manual(caminho_imagem, slug):

    if not os.path.exists(caminho_imagem):

        print("❌ Imagem não encontrada.")
        return IMAGEM_PADRAO

    extensao = Path(caminho_imagem).suffix

    novo_nome = f"{slug}{extensao}"

    destino = os.path.join(
        PASTA_IMG,
        novo_nome
    )

    shutil.copy2(
        caminho_imagem,
        destino
    )

    print(
        f"\n🖼️ Imagem copiada: {novo_nome}"
    )

    return novo_nome

# ==========================================
# TEXTO MANUAL
# ==========================================

def ler_texto_manual():

    print("\nCole a matéria abaixo.")
    print("Digite FIM em uma linha separada.\n")

    linhas = []

    while True:

        linha = input()

        if linha.strip().upper() == "FIM":
            break

        linhas.append(linha)

    return "\n".join(linhas)

# ==========================================
# HTML
# ==========================================

def transformar_em_html(texto):

    html = ""

    paragrafos = texto.split("\n")

    for p in paragrafos:

        p = p.strip()

        if p:

            html += f"<p>{p}</p>\n"

    return html

# ==========================================
# NEWS DATA
# ==========================================

def carregar_news_data():

    if not os.path.exists(
        NEWS_DATA_ARQUIVO
    ):
        return []

    with open(
        NEWS_DATA_ARQUIVO,
        "r",
        encoding="utf-8"
    ) as f:

        conteudo = f.read()

    conteudo = conteudo.replace(
        "window.NEWS_DATA =",
        ""
    ).strip()

    if conteudo.endswith(";"):
        conteudo = conteudo[:-1]

    return json.loads(conteudo)

def salvar_news_data(lista):

    texto = json.dumps(
        lista,
        ensure_ascii=False,
        indent=2
    )

    final = f"window.NEWS_DATA = {texto};"

    with open(
        NEWS_DATA_ARQUIVO,
        "w",
        encoding="utf-8"
    ) as f:

        f.write(final)

# ==========================================
# CRIAR HTML
# ==========================================

def criar_html(
    titulo,
    autor,
    categoria,
    imagem,
    conteudo_html,
    slug,
    relacionadas
):

    data_formatada = datetime.now().strftime(
        "%d/%m/%Y"
    )

    relacionadas_html = ""

    for noticia in relacionadas:

        relacionadas_html += f'''
<li>
  <a href="./{noticia["slug"]}">
    {noticia["title"]}
  </a>
</li>
'''

    html = f'''<!doctype html>
<html lang="pt-BR">

<head>

<meta charset="utf-8">

<meta
  name="viewport"
  content="width=device-width, initial-scale=1"
>

<meta
  name="last-updated"
  content="{datetime.now().strftime("%Y-%m-%d")}"
>

<title>{titulo} — Jornal AJL</title>

<link
  rel="stylesheet"
  href="../CSS/style.css"
>

</head>

<body>

<div id="siteHeader"></div>

<main class="container">

<div class="article-wrap">

<article class="article">

<div class="pad">

<div class="breadcrumb">

<a href="../index.html">
Início
</a>

•

<a href="./todas-noticias.html">
{categoria}
</a>

</div>

<h1>{titulo}</h1>

<div class="byline">

<span>

Por

<strong>

<a
href="https://instagram.com/{autor.replace("@","")}"
target="_blank"
>

{autor}

</a>

</strong>

</span>

<span>•</span>

<span>
Atualizado em {data_formatada}
</span>

</div>

</div>

<div class="media">

<img
src="../IMG/{imagem}"
alt="{titulo}"
>

</div>

<div class="pad body">

{conteudo_html}

<div class="note">

Este é um conteúdo escolar,
produzido por alunos.

</div>

</div>

</article>

<aside class="widget">

<header>

<h3>Leia também</h3>

<a
href="./todas-noticias.html"
class="small"
>

Ver todas

</a>

</header>

<div class="pad">

<ul class="small">

{relacionadas_html}

</ul>

</div>

</aside>

</div>

</main>

<div id="siteFooter"></div>

<script src="../js/site.js"></script>

</body>
</html>
'''

    caminho = os.path.join(
        PASTA_PAGES,
        slug
    )

    with open(
        caminho,
        "w",
        encoding="utf-8"
    ) as f:

        f.write(html)

    print(f"\n✅ HTML criado: {slug}")

# ==========================================
# MAIN
# ==========================================

def main():

    print("\n==============================")
    print("GERADOR DE MATÉRIAS - AJL")
    print("==============================\n")

    titulo = input(
        "Título da matéria: "
    ).strip()

    if not titulo:

        print(
            "❌ Título obrigatório."
        )

        return

    slug_base = slugify(titulo)

    slug = slug_base + ".html"

    caminho_html = os.path.join(
        PASTA_PAGES,
        slug
    )

    if os.path.exists(caminho_html):

        print(
            "❌ Já existe matéria com esse nome."
        )

        return

    autor = input(
        "Instagram do autor (sem @) [opcional]: "
    ).strip()

    if not autor:
        autor = AUTOR_PADRAO
    else:
        autor = (
            "@"
            + autor.replace("@", "")
        )

    print("\nCategorias:")

    for k, v in CATEGORIAS.items():

        print(f"{k} - {v}")

    categoria_input = input(
        "\nEscolha: "
    ).strip()

    categoria = CATEGORIAS.get(
        categoria_input,
        "Cultura"
    )

    caminho_docx = input(
        "\nCole o caminho do .docx ou aperte ENTER para texto manual: "
    ).strip().replace('"', '')

    texto = ""

    imagem = IMAGEM_PADRAO

    if caminho_docx:

        if not os.path.exists(
            caminho_docx
        ):

            print(
                "❌ Arquivo não encontrado."
            )

            return

        texto = ler_docx(
            caminho_docx
        )

        imagem = extrair_imagem_docx(
            caminho_docx,
            slug_base
        )

        # CASO NÃO TENHA IMAGEM NO DOCX
        if imagem == IMAGEM_PADRAO:

            caminho_imagem = input(
                "\nCole o caminho da imagem [opcional]: "
            ).strip().replace('"', '')

            if caminho_imagem:

                imagem = copiar_imagem_manual(
                    caminho_imagem,
                    slug_base
                )

    else:

        imagem_manual = input(
            "\nNome da imagem (ex: foto.jpg): "
        ).strip()

        if imagem_manual:
            imagem = imagem_manual

        texto = ler_texto_manual()

    if not texto:

        print("❌ Texto vazio.")
        return

    excerpt_manual = input(
        "\nResumo manual [opcional]: "
    ).strip()

    excerpt = (
        excerpt_manual
        if excerpt_manual
        else gerar_excerpt(texto)
    )

    conteudo_html = transformar_em_html(
        texto
    )

    noticias = carregar_news_data()

    relacionadas = noticias[-6:]

    nova_noticia = {
        "title": titulo,
        "slug": slug,
        "image": imagem,
        "excerpt": excerpt,
        "author": autor,
        "category": categoria
    }

    noticias.append(nova_noticia)

    salvar_news_data(
        noticias
    )

    criar_html(
        titulo,
        autor,
        categoria,
        imagem,
        conteudo_html,
        slug,
        relacionadas
    )

    print("\n✅ NEWS_DATA atualizado.")
    print("✅ Processo finalizado.")

    print("\nAgora é só:")

    print("git add .")
    print('git commit -m "Nova matéria"')
    print("git push")

if __name__ == "__main__":
    main()